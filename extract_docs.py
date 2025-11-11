import PyPDF2
from docx import Document

# Extract PDF content
print("=" * 80)
print("EXTRACTING PDF: GuardElns Synopsis.pdf")
print("=" * 80)
try:
    with open('GuardElns Synopsis.pdf', 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        print(f"\nTotal pages: {len(pdf_reader.pages)}\n")
        for i, page in enumerate(pdf_reader.pages):
            print(f"\n--- Page {i+1} ---\n")
            print(page.extract_text())
except Exception as e:
    print(f"Error reading PDF: {e}")

print("\n\n")
print("=" * 80)
print("EXTRACTING DOCX: Final Project Report Format CSE (1).docx")
print("=" * 80)
try:
    doc = Document('Final Project Report Format CSE (1).docx')
    print(f"\nTotal paragraphs: {len(doc.paragraphs)}\n")
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
except Exception as e:
    print(f"Error reading DOCX: {e}")
